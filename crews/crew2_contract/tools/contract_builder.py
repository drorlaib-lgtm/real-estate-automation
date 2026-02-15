"""Contract Builder Tool - Fills the reca agreement.docx template with transaction data."""

import os
import json
import re
import copy
from pathlib import Path

from docx import Document
from crewai.tools import tool

# Path to the contract template
TEMPLATE_PATH = Path(__file__).parent.parent.parent.parent / "templates" / "reca_agreement.docx"


def number_to_hebrew_words(n: int) -> str:
    """Convert number to Hebrew words for contract amounts."""
    if n == 0:
        return "אפס ש\"ח"

    units_w = {
        1: "אחד", 2: "שניים", 3: "שלושה", 4: "ארבעה", 5: "חמישה",
        6: "שישה", 7: "שבעה", 8: "שמונה", 9: "תשעה",
    }
    tens_w = {
        10: "עשרה", 20: "עשרים", 30: "שלושים", 40: "ארבעים", 50: "חמישים",
        60: "שישים", 70: "שבעים", 80: "שמונים", 90: "תשעים",
    }
    hundred_w = {
        100: "מאה", 200: "מאתיים", 300: "שלוש מאות", 400: "ארבע מאות",
        500: "חמש מאות", 600: "שש מאות", 700: "שבע מאות",
        800: "שמונה מאות", 900: "תשע מאות",
    }

    def _small(num):
        """Convert 1-999 to Hebrew words."""
        if num == 0:
            return ""
        result = []
        if num >= 100:
            h = (num // 100) * 100
            result.append(hundred_w[h])
            num %= 100
        if num >= 10:
            t = (num // 10) * 10
            u = num % 10
            if u:
                result.append(f"{tens_w[t]} ו{units_w[u]}")
            else:
                result.append(tens_w[t])
        elif num > 0:
            result.append(units_w[num])
        return " ".join(result)

    parts = []

    # Millions
    if n >= 1_000_000:
        m = n // 1_000_000
        n %= 1_000_000
        if m == 1:
            parts.append("מיליון")
        elif m == 2:
            parts.append("שני מיליון")
        else:
            parts.append(f"{_small(m)} מיליון")

    # Thousands
    if n >= 1000:
        t = n // 1000
        n %= 1000
        if t == 1:
            parts.append("אלף")
        elif t == 2:
            parts.append("אלפיים")
        else:
            parts.append(f"{_small(t)} אלפים")

    # Hundreds/tens/units
    rest = _small(n)
    if rest:
        parts.append(rest)

    return " ".join(parts) + " ש\"ח"


def format_price_hebrew(price) -> str:
    """Format price: 2,500,000 ₪ (שני מיליון וחמש מאות אלף ש\"ח)."""
    price = int(float(price)) if price else 0
    return f"{price:,} ₪ ({number_to_hebrew_words(price)})"


def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace placeholders in a paragraph, handling split runs.

    DOCX often splits placeholder text across multiple runs (e.g. '{{', 'FIELD', '}}').
    This function joins all runs, performs replacements, and redistributes.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = full_text
    changed = False
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value))
            changed = True

    if changed and paragraph.runs:
        # Put the full replaced text in the first run, clear the rest
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def build_replacements(data: dict) -> dict:
    """Build the placeholder -> value mapping from transaction data."""
    price = int(float(data.get("price", 0) or 0))

    # Payment schedule: default 10% / tax advance / remainder
    payment_1 = data.get("payment_1", "")
    payment_2 = data.get("payment_2", "")
    payment_3 = data.get("payment_3", "")

    if not payment_1 and price:
        payment_1 = f"{int(price * 0.10):,} ₪"
    if not payment_2 and price:
        payment_2 = f"{int(price * 0.15):,} ₪"
    if not payment_3 and price:
        payment_3 = f"{int(price - int(price * 0.10) - int(price * 0.15)):,} ₪"

    escrow_amount = data.get("escrow_amount", "")
    if not escrow_amount and price:
        escrow_amount = f"{int(price * 0.10):,}"

    # Storage text
    storage_val = data.get("storage", "no")
    storage_text = "מחסן" if storage_val == "yes" else ""

    # Parking text
    parking_map = {
        "none": "ללא",
        "covered": "מקורה",
        "uncovered": "לא מקורה",
        "underground": "תת-קרקעית"
    }
    parking_val = data.get("parking", "none")
    parking_text = parking_map.get(parking_val, "ללא")

    # Seller notes
    seller_notes = (
        data.get("seller_declaration_notes", "")
        or data.get("notes", "")
        or ""
    )

    replacements = {
        # Well-formed placeholders
        "{{SIGNING_DATE}}": data.get("signing_date", "________"),
        "{{SELLER_NAME}}": data.get("seller_name", "________"),
        "{{SELLER_ID}}": data.get("seller_id", "________"),
        "{{SELLER_ADDRESS}}": data.get("seller_address", "________"),
        "{{SELLER_PHONE}}": data.get("seller_phone", "________"),
        "{{SELLER_EMAIL}}": data.get("seller_email", "________"),
        "{{SELLER2_NAME}}": data.get("seller2_name", "________"),
        "{{SELLER2_ID}}": data.get("seller2_id", "________"),
        "{{BUYER_NAME}}": data.get("buyer_name", "________"),
        "{{BUYER_ID}}": data.get("buyer_id", "________"),
        "{{BUYER_ADDRESS}}": data.get("buyer_address", "________"),
        "{{BUYER_PHONE}}": data.get("buyer_phone", "________"),
        "{{BUYER_EMAIL}}": data.get("buyer_email", "________"),
        "{{BUYER2_NAME}}": data.get("buyer2_name", "________"),
        "{{BUYER2_ID}}": data.get("buyer2_id", "________"),
        "{{ROOMS}}": str(data.get("rooms", "____")),
        "{{FLOOR}}": str(data.get("floor", "____")),
        "{{PROPERTY_ADDRESS}}": data.get("property_address", "________"),
        "{{STORAGE}}": storage_text,
        "{{PARKING}}": parking_text,
        "{{BLOCK}}": str(data.get("block_number", "____")),
        "{{PARCEL}}": str(data.get("parcel_number", "____")),
        "{{SUB_PARCEL}}": str(data.get("sub_parcel", "____")),
        "{{DELIVERY_DATE}}": data.get("delivery_date", "________"),
        "{{SELLER_NOTES}}": seller_notes,
        "{{PRICE}}": f"{price:,}",
        "{{PRICE_WORDS}}": number_to_hebrew_words(price),
        "{{PAYMENT_1}}": str(payment_1),
        "{{PAYMENT_2}}": str(payment_2),
        "{{PAYMENT_3}}": str(payment_3),
        "{{ESCROW_AMOUNT}}": str(escrow_amount),
        "{{BUYER_LAWYER}}": data.get("buyer_lawyer", "________"),
        "{{BUYER_LAWYER_EMAIL}}": data.get("buyer_lawyer_email", "________"),
        "{{MORTGAGE_BANK}}": data.get("mortgage_bank", "________"),

        # Malformed placeholders found in the template
        "{{ SELLER _EMAIL}": data.get("seller_email", "________"),
        "{{BUYER_EMAIL}": data.get("buyer_email", "________"),
        "{PRICE}}": f"{price:,}",
    }

    return replacements


def build_contract_document(data: dict, variation: str = "standard") -> Document:
    """Fill the reca agreement.docx template with transaction data.

    Args:
        data: Flat transaction data dict.
        variation: Ignored (kept for API compatibility). Template is always the same.

    Returns:
        Filled Document object.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"תבנית חוזה לא נמצאה: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))
    replacements = build_replacements(data)

    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Replace in tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # Replace in headers/footers
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    return doc


@tool("build_contract")
def build_contract(clean_data_json: str, variation: str = "standard") -> str:
    """Build a Hebrew real estate contract by filling the template.
    Returns path to the generated contract file."""
    data = json.loads(clean_data_json)
    doc = build_contract_document(data, variation)

    os.makedirs("artifacts", exist_ok=True)
    filename = f"artifacts/contract_{variation}.docx"
    doc.save(filename)
    return f"חוזה נוצר בהצלחה: {filename}"
