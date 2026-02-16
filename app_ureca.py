"""
URECA NOW - Client Portal for Real Estate Transactions
=======================================================
Simplified client-facing app: data entry + document upload.
All processing (validation, contract, Drive, email) happens behind the scenes.
"""

import os
import sys
import json
import io
import streamlit as st
from pathlib import Path
from datetime import date, datetime, timedelta

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document as DocxDocument
from crews.crew1_data.tools.validator import run_validation
from crews.crew1_data.tools.data_cleaner import merge_and_clean
from crews.crew2_contract.tools.contract_builder import (
    build_replacements, _replace_in_paragraph,
)
from tools.data_adapter import normalize_transaction
from tools.submission_manager import save_submission, upload_submission_to_drive
from email_service import send_notification_email

# Template for ureca - RECA agreement (the only correct template)
TEMPLATE_PATH = Path(__file__).parent / "templates" / "reca_agreement.docx"

# European date format
EU_DATE_FMT = "%d/%m/%Y"


def parse_date_str(date_str: str, fallback: date = None) -> date:
    """Parse date string in DD/MM/YYYY or YYYY-MM-DD format."""
    if not date_str:
        return fallback
    for fmt in (EU_DATE_FMT, "%Y-%m-%d"):
        try:
            return datetime.strptime(date_str, fmt).date()
        except (ValueError, TypeError):
            continue
    return fallback


# Page config
st.set_page_config(
    page_title="URECA NOW",
    page_icon="🏠",
    layout="centered",
)

# Viewport meta tag for mobile - prevents auto zoom
st.markdown("""
<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
""", unsafe_allow_html=True)

# Light blue theme CSS
st.markdown("""
<style>
    /* Light blue background */
    .stApp {
        direction: rtl;
        background-color: #A8D4E6;
        color: #1B2A4A;
    }

    /* Main container */
    .main .block-container {
        direction: rtl;
        text-align: right;
    }

    /* Text elements */
    .stMarkdown, .stMarkdown p, .stMarkdown li {
        direction: rtl;
        text-align: right;
        color: #1B2A4A !important;
    }

    /* Headers */
    h1, h2, h3 {
        text-align: center;
        color: #1B2A4A !important;
    }

    /* Title styling */
    h1 {
        font-size: 2.5rem !important;
        font-weight: 800 !important;
        letter-spacing: 3px;
        text-shadow: 0 2px 10px rgba(27, 42, 74, 0.2);
    }

    /* Labels */
    .stTextInput label, .stSelectbox label, .stDateInput label,
    .stNumberInput label, .stTextArea label, .stFileUploader label {
        text-align: right;
        color: #2C3E50 !important;
    }

    /* Input fields */
    .stTextInput input, .stNumberInput input, .stTextArea textarea {
        direction: rtl;
        text-align: right;
        background-color: #FFFFFF !important;
        color: #1B2A4A !important;
        border: 1px solid #7AB8D4 !important;
        border-radius: 8px !important;
    }

    .stTextInput input:focus, .stNumberInput input:focus, .stTextArea textarea:focus {
        border-color: #2980B9 !important;
        box-shadow: 0 0 8px rgba(41, 128, 185, 0.4) !important;
    }

    /* Select boxes */
    .stSelectbox > div > div {
        background-color: #FFFFFF !important;
        color: #1B2A4A !important;
        border: 1px solid #7AB8D4 !important;
        border-radius: 8px !important;
    }

    /* Date input */
    .stDateInput > div > div > input {
        background-color: #FFFFFF !important;
        color: #1B2A4A !important;
        border: 1px solid #7AB8D4 !important;
    }

    /* Buttons */
    .stButton > button {
        background-color: #2980B9 !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        font-size: 1.1rem !important;
        font-weight: 600 !important;
        padding: 0.6rem 2rem !important;
        transition: all 0.3s ease !important;
    }

    .stButton > button:hover {
        background-color: #1F6FA0 !important;
        box-shadow: 0 4px 15px rgba(41, 128, 185, 0.4) !important;
        transform: translateY(-1px) !important;
    }

    .stButton > button[kind="primary"] {
        background-color: #10B981 !important;
        font-size: 1.3rem !important;
        padding: 0.8rem 3rem !important;
    }

    .stButton > button[kind="primary"]:hover {
        background-color: #059669 !important;
        box-shadow: 0 4px 20px rgba(16, 185, 129, 0.4) !important;
    }

    /* Expanders */
    .streamlit-expanderHeader {
        background-color: #8FC4D8 !important;
        color: #1B2A4A !important;
        border-radius: 8px !important;
    }

    [data-testid="stExpander"] {
        background-color: #8FC4D8 !important;
        border: 1px solid #7AB8D4 !important;
        border-radius: 8px !important;
    }

    [data-testid="stExpander"] details {
        background-color: #8FC4D8 !important;
    }

    [data-testid="stExpander"] summary {
        background-color: #8FC4D8 !important;
        color: #1B2A4A !important;
    }

    [data-testid="stExpander"] details summary span {
        color: #1B2A4A !important;
    }

    [data-testid="stExpander"] summary:hover {
        background-color: #7BB5CE !important;
    }

    [data-testid="stExpanderDetails"] {
        background-color: #96C8DB !important;
    }

    /* Expander toggle icon */
    [data-testid="stExpander"] summary svg {
        fill: #1B2A4A !important;
        color: #1B2A4A !important;
    }

    /* File uploader */
    [data-testid="stFileUploader"] {
        border: 2px dashed #7AB8D4 !important;
        border-radius: 10px !important;
        padding: 20px !important;
        text-align: center !important;
        background-color: #96C8DB !important;
    }

    [data-testid="stFileUploader"]:hover {
        border-color: #2980B9 !important;
        background-color: #8FC4D8 !important;
    }

    /* File uploader - file name text */
    [data-testid="stFileUploader"] small,
    [data-testid="stFileUploader"] span,
    [data-testid="stFileUploader"] div,
    [data-testid="stFileUploader"] p,
    [data-testid="stFileUploader"] [data-testid="stMarkdownContainer"],
    [data-testid="stFileUploader"] [data-testid="stMarkdownContainer"] p {
        color: #1B2A4A !important;
    }

    [data-testid="stFileUploaderFileName"] {
        color: #1B2A4A !important;
    }

    [data-testid="stFileUploader"] section > div {
        color: #1B2A4A !important;
    }

    [data-testid="stFileUploader"] button {
        color: #1B2A4A !important;
    }

    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #1B2A4A !important;
    }

    [data-testid="stSidebar"] .stMarkdown p,
    [data-testid="stSidebar"] .stMarkdown {
        color: #D0E8F2 !important;
    }

    [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
        color: #FFFFFF !important;
    }

    /* Progress bar */
    .stProgress > div > div {
        direction: ltr;
        background-color: #2980B9 !important;
    }

    /* Success/info/warning messages */
    .stSuccess, .stInfo, .stWarning {
        border-radius: 8px !important;
    }

    /* Dividers */
    hr {
        border-color: #7AB8D4 !important;
    }

    /* Hide Streamlit branding and Deploy button */
    #MainMenu, footer, [data-testid="stAppDeployButton"] { display: none !important; }

    /* Hide top white header bar */
    header[data-testid="stHeader"] {
        background-color: #A8D4E6 !important;
        height: 0px !important;
        min-height: 0px !important;
        padding: 0 !important;
    }

    /* Sticky header */
    .sticky-header {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        z-index: 999;
        background-color: #A8D4E6;
        direction: ltr !important;
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 12px 0;
        border-bottom: 2px solid #7AB8D4;
    }

    .sticky-header-inner {
        display: flex;
        direction: ltr !important;
        align-items: center;
        justify-content: center;
        gap: 50px;
        width: 100%;
        max-width: 600px;
        margin: 0 auto;
        transform: translateX(-80px);
    }

    .sticky-header img {
        height: 55px;
        width: auto;
    }

    .sticky-header h1 {
        margin: 0;
        font-size: 2.2rem !important;
        font-weight: 800 !important;
        letter-spacing: 3px;
        color: #1B2A4A !important;
        text-shadow: 0 2px 10px rgba(27, 42, 74, 0.2);
        white-space: nowrap;
    }

    /* Push content below sticky header */
    .main .block-container {
        padding-top: 100px !important;
    }

    /* Horizontal blocks RTL */
    [data-testid="stHorizontalBlock"] { flex-direction: row-reverse; }

    /* =================== MOBILE RESPONSIVE =================== */
    @media (max-width: 768px) {
        /* Sticky header - smaller on mobile */
        .sticky-header {
            padding: 8px 0;
        }

        .sticky-header-inner {
            gap: 12px;
            max-width: 100%;
            padding: 0 10px;
            transform: none !important;
        }

        .sticky-header img {
            height: 32px;
        }

        .sticky-header h1 {
            font-size: 1.2rem !important;
            letter-spacing: 1px;
        }

        /* Push content below smaller header */
        .main .block-container {
            padding-top: 70px !important;
            padding-left: 8px !important;
            padding-right: 8px !important;
        }

        /* Stack columns vertically on mobile */
        [data-testid="stHorizontalBlock"] {
            flex-direction: column !important;
            gap: 0 !important;
        }

        [data-testid="stHorizontalBlock"] > [data-testid="stColumn"] {
            width: 100% !important;
            flex: 1 1 100% !important;
            min-width: 100% !important;
        }

        /* Smaller headers */
        h1 {
            font-size: 1.5rem !important;
        }

        h2 {
            font-size: 1.2rem !important;
        }

        h3 {
            font-size: 1rem !important;
        }

        /* Bigger touch targets for inputs */
        .stTextInput input,
        .stNumberInput input,
        .stTextArea textarea {
            font-size: 16px !important;
            padding: 10px !important;
            min-height: 44px !important;
        }

        .stSelectbox > div > div {
            min-height: 44px !important;
            font-size: 16px !important;
        }

        .stDateInput > div > div > input {
            font-size: 16px !important;
            min-height: 44px !important;
        }

        /* Bigger buttons on mobile */
        .stButton > button {
            font-size: 1rem !important;
            padding: 12px 16px !important;
            min-height: 48px !important;
            width: 100% !important;
        }

        .stButton > button[kind="primary"] {
            font-size: 1.1rem !important;
            padding: 14px 16px !important;
        }

        /* File uploader - bigger area */
        [data-testid="stFileUploader"] {
            padding: 14px !important;
        }

        /* Expanders - easier to tap */
        [data-testid="stExpander"] summary {
            padding: 12px !important;
            min-height: 48px !important;
        }

        /* Labels - readable on mobile */
        .stTextInput label,
        .stSelectbox label,
        .stDateInput label,
        .stNumberInput label,
        .stTextArea label,
        .stFileUploader label {
            font-size: 14px !important;
        }

        /* Hide sidebar by default on mobile */
        [data-testid="stSidebar"] {
            display: none !important;
        }

        /* Progress bar full width */
        .stProgress {
            width: 100% !important;
        }
    }

    /* Extra small screens (phones in portrait) */
    @media (max-width: 480px) {
        .sticky-header-inner {
            gap: 8px;
        }

        .sticky-header img {
            height: 26px;
        }

        .sticky-header h1 {
            font-size: 1rem !important;
            letter-spacing: 0.5px;
        }

        .main .block-container {
            padding-top: 60px !important;
            padding-left: 4px !important;
            padding-right: 4px !important;
        }

        h1 {
            font-size: 1.3rem !important;
        }

        h2 {
            font-size: 1.1rem !important;
        }
    }

    /* Viewport meta for proper mobile scaling */
    @media screen and (max-width: 768px) {
        .stApp {
            -webkit-text-size-adjust: 100%;
        }
    }
</style>
""", unsafe_allow_html=True)

# Sticky header with logos
import base64

def img_to_base64(path):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

logo_dror_b64 = img_to_base64("assets/logo_dror.png")
logo_meir_b64 = img_to_base64("assets/logo_meir.jpeg")

st.markdown(f"""
<div class="sticky-header">
    <div class="sticky-header-inner">
        <img src="data:image/jpeg;base64,{logo_meir_b64}" alt="logo meir">
        <h1>URECA NOW</h1>
        <img src="data:image/png;base64,{logo_dror_b64}" alt="logo dror">
    </div>
</div>
""", unsafe_allow_html=True)
st.markdown("---")

# Sidebar - step indicator
st.sidebar.header("שלבים")
pages = ["📝 הזנת נתונים", "📄 העלאת מסמכים"]

if "current_page" not in st.session_state:
    st.session_state.current_page = 0

page = pages[st.session_state.current_page]
st.sidebar.markdown(f"**שלב נוכחי:** {page}")

for i, p in enumerate(pages):
    if i < st.session_state.current_page:
        st.sidebar.markdown(f"✅ {p}")
    elif i == st.session_state.current_page:
        st.sidebar.markdown(f"➡️ **{p}**")
    else:
        st.sidebar.markdown(f"⬜ {p}")

# Initialize session state
if "sellers" not in st.session_state:
    st.session_state.sellers = [{"name": "", "id": "", "address": "", "phone": "", "email": "", "marital_status": ""}]
if "buyers" not in st.session_state:
    st.session_state.buyers = [{"name": "", "id": "", "address": "", "phone": "", "email": "", "marital_status": ""}]
if "property_data" not in st.session_state:
    st.session_state.property_data = {}
if "transaction_data" not in st.session_state:
    st.session_state.transaction_data = {}
if "seller_notes" not in st.session_state:
    st.session_state.seller_notes = ""
if "uploaded_docs" not in st.session_state:
    st.session_state.uploaded_docs = {}
if "sent" not in st.session_state:
    st.session_state.sent = False

MARITAL_OPTIONS = ["", "רווק/ה", "נשוי/אה", "גרוש/ה", "אלמן/ה"]
MARITAL_MAP = {"": "", "רווק/ה": "single", "נשוי/אה": "married", "גרוש/ה": "divorced", "אלמן/ה": "widowed"}
MARITAL_MAP_REVERSE = {v: k for k, v in MARITAL_MAP.items()}


def render_person_form(person_type: str, index: int, data: dict) -> dict:
    """Render form fields for a person."""
    prefix = f"{person_type}_{index}"
    col1, col2 = st.columns(2)
    with col1:
        data["id"] = st.text_input("תעודת זהות *", value=data.get("id", ""), key=f"{prefix}_id")
        data["email"] = st.text_input("דוא\"ל *", value=data.get("email", ""), key=f"{prefix}_email")
    with col2:
        data["name"] = st.text_input("שם מלא *", value=data.get("name", ""), key=f"{prefix}_name")
        data["phone"] = st.text_input("טלפון *", value=data.get("phone", ""), placeholder="05XXXXXXXX", key=f"{prefix}_phone")
        data["address"] = st.text_input("כתובת *", value=data.get("address", ""), key=f"{prefix}_address")
        current_marital = MARITAL_MAP_REVERSE.get(data.get("marital_status", ""), "")
        marital_index = MARITAL_OPTIONS.index(current_marital) if current_marital in MARITAL_OPTIONS else 0
        marital = st.selectbox("מצב משפחתי *", MARITAL_OPTIONS, index=marital_index, key=f"{prefix}_marital")
        data["marital_status"] = MARITAL_MAP.get(marital, "")
    return data


def build_client_data() -> dict:
    """Build flat client_data from session state."""
    primary_seller = st.session_state.sellers[0] if st.session_state.sellers else {}
    secondary_seller = st.session_state.sellers[1] if len(st.session_state.sellers) > 1 else {}
    primary_buyer = st.session_state.buyers[0] if st.session_state.buyers else {}
    secondary_buyer = st.session_state.buyers[1] if len(st.session_state.buyers) > 1 else {}
    return {
        "seller_name": primary_seller.get("name", ""),
        "seller_id": primary_seller.get("id", ""),
        "seller_address": primary_seller.get("address", ""),
        "seller_phone": primary_seller.get("phone", ""),
        "seller_email": primary_seller.get("email", ""),
        "seller_marital_status": primary_seller.get("marital_status", ""),
        "seller2_name": secondary_seller.get("name", ""),
        "seller2_id": secondary_seller.get("id", ""),
        "buyer_name": primary_buyer.get("name", ""),
        "buyer_id": primary_buyer.get("id", ""),
        "buyer_address": primary_buyer.get("address", ""),
        "buyer_phone": primary_buyer.get("phone", ""),
        "buyer_email": primary_buyer.get("email", ""),
        "buyer_marital_status": primary_buyer.get("marital_status", ""),
        "buyer2_name": secondary_buyer.get("name", ""),
        "buyer2_id": secondary_buyer.get("id", ""),
        "property_address": st.session_state.property_data.get("address", ""),
        "block_number": st.session_state.property_data.get("block_number", ""),
        "parcel_number": st.session_state.property_data.get("parcel_number", ""),
        "sub_parcel": st.session_state.property_data.get("sub_parcel", ""),
        "area_sqm": str(st.session_state.property_data.get("area_sqm", "")),
        "rooms": str(st.session_state.property_data.get("rooms", "")),
        "floor": str(st.session_state.property_data.get("floor", "")),
        "property_type": st.session_state.property_data.get("property_type", "apartment"),
        "parking": st.session_state.property_data.get("parking", "none"),
        "storage": st.session_state.property_data.get("storage", "no"),
        "price": str(st.session_state.transaction_data.get("price", "")),
        "signing_date": st.session_state.transaction_data.get("signing_date", ""),
        "delivery_date": st.session_state.transaction_data.get("delivery_date", ""),
        "notes": st.session_state.seller_notes,
        "seller_declaration_notes": st.session_state.seller_notes,
        "mortgage_bank": st.session_state.transaction_data.get("mortgage_bank", ""),
        "buyer_lawyer": st.session_state.transaction_data.get("buyer_lawyer", ""),
        "buyer_lawyer_email": st.session_state.transaction_data.get("buyer_lawyer_email", ""),
        "all_sellers": st.session_state.sellers,
        "all_buyers": st.session_state.buyers,
    }


def build_contract_from_template(data: dict) -> DocxDocument:
    """Fill agreement X.10 template with transaction data."""
    doc = DocxDocument(str(TEMPLATE_PATH))
    replacements = build_replacements(data)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for paragraph in hf.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    return doc


# =================== PAGE 1: DATA ENTRY ===================
if page == "📝 הזנת נתונים":
    st.header("הזנת פרטי העסקה")

    # Sellers
    st.subheader("פרטי המוכרים")
    for i in range(len(st.session_state.sellers)):
        with st.expander(
            f"מוכר {i + 1}" + (f" - {st.session_state.sellers[i].get('name', '')}" if st.session_state.sellers[i].get('name') else ""),
            expanded=(i == 0),
        ):
            st.session_state.sellers[i] = render_person_form("seller", i, st.session_state.sellers[i])
            if len(st.session_state.sellers) > 1:
                if st.button(f"הסר מוכר {i + 1}", key=f"remove_seller_{i}"):
                    st.session_state.sellers.pop(i)
                    st.rerun()

    if st.button("+ הוסף מוכר נוסף", key="add_seller"):
        st.session_state.sellers.append({"name": "", "id": "", "address": "", "phone": "", "email": "", "marital_status": ""})
        st.rerun()

    st.session_state.transaction_data["mortgage_bank"] = st.text_input(
        "משכנתא בבנק",
        value=st.session_state.transaction_data.get("mortgage_bank", ""),
        key="mortgage_bank_input",
        placeholder="שם הבנק (אם קיימת משכנתא)",
    )

    st.markdown("---")

    # Buyers
    st.subheader("פרטי הקונים")
    for i in range(len(st.session_state.buyers)):
        with st.expander(
            f"קונה {i + 1}" + (f" - {st.session_state.buyers[i].get('name', '')}" if st.session_state.buyers[i].get('name') else ""),
            expanded=(i == 0),
        ):
            st.session_state.buyers[i] = render_person_form("buyer", i, st.session_state.buyers[i])
            if len(st.session_state.buyers) > 1:
                if st.button(f"הסר קונה {i + 1}", key=f"remove_buyer_{i}"):
                    st.session_state.buyers.pop(i)
                    st.rerun()

    if st.button("+ הוסף קונה נוסף", key="add_buyer"):
        st.session_state.buyers.append({"name": "", "id": "", "address": "", "phone": "", "email": "", "marital_status": ""})
        st.rerun()

    col_bl1, col_bl2 = st.columns(2)
    with col_bl1:
        st.session_state.transaction_data["buyer_lawyer_email"] = st.text_input(
            "מייל עו\"ד הקונה",
            value=st.session_state.transaction_data.get("buyer_lawyer_email", ""),
            key="buyer_lawyer_email_input",
            placeholder="lawyer@example.com",
        )
    with col_bl2:
        st.session_state.transaction_data["buyer_lawyer"] = st.text_input(
            "שם עו\"ד הקונה",
            value=st.session_state.transaction_data.get("buyer_lawyer", ""),
            key="buyer_lawyer_input",
            placeholder="שם עורך הדין של הקונה",
        )

    st.markdown("---")

    # Property
    st.subheader("פרטי הנכס")
    prop = st.session_state.property_data

    col3, col4 = st.columns(2)
    with col3:
        prop["area_sqm"] = st.number_input("שטח (מ\"ר) *", min_value=10, max_value=5000, value=int(prop.get("area_sqm", 80)), key="area")
        prop["rooms"] = st.number_input("חדרים *", min_value=1.0, max_value=20.0, value=float(prop.get("rooms", 3.0)), step=0.5, key="rooms_input")
        prop["floor"] = st.number_input("קומה", min_value=-2, max_value=100, value=int(prop.get("floor", 0)), key="floor_input")

        prop_types = ["דירה", "פנטהאוז", "דירת גן", "דופלקס", "בית פרטי", "מגרש"]
        type_map = {"דירה": "apartment", "פנטהאוז": "penthouse", "דירת גן": "garden", "דופלקס": "duplex", "בית פרטי": "house", "מגרש": "land"}
        type_map_reverse = {v: k for k, v in type_map.items()}
        current_type = type_map_reverse.get(prop.get("property_type", "apartment"), "דירה")
        prop_type = st.selectbox("סוג נכס *", prop_types, index=prop_types.index(current_type), key="prop_type")
        prop["property_type"] = type_map.get(prop_type, "apartment")
    with col4:
        prop["address"] = st.text_input("כתובת הנכס *", value=prop.get("address", ""), key="prop_addr")
        prop["block_number"] = st.text_input("גוש *", value=prop.get("block_number", ""), key="block")
        prop["parcel_number"] = st.text_input("חלקה *", value=prop.get("parcel_number", ""), key="parcel")
        prop["sub_parcel"] = st.text_input("תת-חלקה", value=prop.get("sub_parcel", ""), key="sub_parcel")

    col5, col6 = st.columns(2)
    with col5:
        storage_options = ["לא", "כן"]
        current_storage = "כן" if prop.get("storage") == "yes" else "לא"
        storage = st.selectbox("מחסן", storage_options, index=storage_options.index(current_storage), key="storage_input")
        prop["storage"] = "yes" if storage == "כן" else "no"
    with col6:
        parking_options = ["ללא", "מקורה", "לא מקורה", "תת-קרקעית"]
        parking_map = {"ללא": "none", "מקורה": "covered", "לא מקורה": "uncovered", "תת-קרקעית": "underground"}
        parking_map_reverse = {v: k for k, v in parking_map.items()}
        current_parking = parking_map_reverse.get(prop.get("parking", "none"), "ללא")
        parking = st.selectbox("חניה", parking_options, index=parking_options.index(current_parking), key="parking_input")
        prop["parking"] = parking_map.get(parking, "none")

    st.session_state.property_data = prop

    st.markdown("---")

    # Transaction
    st.subheader("פרטי העסקה")
    trans = st.session_state.transaction_data

    col7, col8 = st.columns(2)
    with col7:
        default_del_date = parse_date_str(trans.get("delivery_date", ""), date.today() + timedelta(days=90))
        trans["delivery_date"] = st.date_input("תאריך מסירה *", value=default_del_date, key="del_date", format="DD/MM/YYYY").strftime(EU_DATE_FMT)
        default_sign_date = parse_date_str(trans.get("signing_date", ""), date.today() + timedelta(days=7))
        trans["signing_date"] = st.date_input("תאריך חתימה *", value=default_sign_date, key="sign_date", format="DD/MM/YYYY").strftime(EU_DATE_FMT)
    with col8:
        trans["price"] = st.number_input("מחיר (₪) *", min_value=50000, max_value=100000000, value=int(trans.get("price", 1500000)), step=50000, key="price_input")

    st.session_state.transaction_data = trans

    st.markdown("---")

    # Seller notes
    st.subheader("הערות המוכר")
    st.session_state.seller_notes = st.text_area(
        "הערות נוספות של המוכר לגבי הנכס",
        value=st.session_state.seller_notes,
        key="seller_notes_input",
        height=150,
    )

    st.markdown("---")

    # Save button -> auto-navigate to documents page
    if st.button("שמור נתונים והמשך", type="primary", use_container_width=True):
        st.session_state.current_page = 1
        st.rerun()


# =================== PAGE 2: DOCUMENT UPLOAD ===================
elif page == "📄 העלאת מסמכים":
    st.header("📄 העלאת מסמכים")
    st.markdown("ניתן לגרור קבצים לאזור ההעלאה או ללחוץ לבחירה")
    st.markdown("---")

    st.subheader("מסמכים נדרשים")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**תעודות זהות מוכרים**")
        seller_ids = st.file_uploader("העלה ת.ז. מוכרים", type=None, accept_multiple_files=True, key="seller_ids")
        if seller_ids:
            st.session_state.uploaded_docs["seller_ids"] = seller_ids

    with col2:
        st.markdown("**תעודות זהות קונים**")
        buyer_ids = st.file_uploader("העלה ת.ז. קונים", type=None, accept_multiple_files=True, key="buyer_ids")
        if buyer_ids:
            st.session_state.uploaded_docs["buyer_ids"] = buyer_ids

    st.markdown("---")

    st.subheader("מסמכי נכס")
    col3, col4 = st.columns(2)
    with col3:
        st.markdown("**נסח טאבו**")
        tabu = st.file_uploader("נסח רישום מקרקעין", type=None, key="tabu")
        if tabu:
            st.session_state.uploaded_docs["tabu"] = tabu
    with col4:
        st.markdown("**אישור עירייה / ועד בית**")
        municipal = st.file_uploader("מסמך עירייה או ועד בית", type=None, key="municipal")
        if municipal:
            st.session_state.uploaded_docs["municipal"] = municipal

    col5, col6 = st.columns(2)
    with col5:
        st.markdown("**אישור ארנונה**")
        arnona = st.file_uploader("אישור תשלום ארנונה", type=None, key="arnona")
        if arnona:
            st.session_state.uploaded_docs["arnona"] = arnona
    with col6:
        st.markdown("**חוזה רכישה קודם (אופציונלי)**")
        purchase = st.file_uploader("חוזה רכישה מקורי של המוכר", type=None, key="purchase")
        if purchase:
            st.session_state.uploaded_docs["purchase"] = purchase

    st.markdown("---")
    st.subheader("מסמכים נוספים")
    other = st.file_uploader("העלה מסמכים נוספים", type=None, accept_multiple_files=True, key="other")
    if other:
        st.session_state.uploaded_docs["other"] = other

    # Summary
    st.markdown("---")
    st.subheader("סיכום מסמכים")
    docs = st.session_state.uploaded_docs
    if docs.get("seller_ids"):
        st.markdown(f"✅ תעודות זהות מוכרים: {len(docs['seller_ids'])}")
    if docs.get("buyer_ids"):
        st.markdown(f"✅ תעודות זהות קונים: {len(docs['buyer_ids'])}")
    if docs.get("tabu"):
        st.markdown("✅ נסח טאבו")
    if docs.get("municipal"):
        st.markdown("✅ מסמך עירייה")
    if docs.get("arnona"):
        st.markdown("✅ אישור ארנונה")
    if docs.get("purchase"):
        st.markdown("✅ חוזה רכישה קודם")
    if docs.get("other"):
        st.markdown(f"✅ מסמכים נוספים: {len(docs['other'])}")

    st.markdown("---")

    # Navigation
    col_nav1, col_nav2 = st.columns([1, 1])
    with col_nav1:
        if st.button("חזור", key="back", use_container_width=True):
            st.session_state.current_page = 0
            st.rerun()
    with col_nav2:
        if st.button("שלח לעורך הדין שלי", type="primary", key="send", use_container_width=True):
            progress = st.progress(0)
            status = st.empty()

            try:
                # Build client data
                client_data = build_client_data()
                seller_name = st.session_state.sellers[0].get("name", "לא ידוע")
                property_address = st.session_state.property_data.get("address", "לא ידוע")

                # Step 1: Validation and generate report
                status.text("מעבד נתונים...")
                progress.progress(15)
                os.makedirs("artifacts", exist_ok=True)

                # Run validation in background (no blocking)
                validation = run_validation(client_data)
                from crews.crew1_data.tools.validator import generate_eda_report
                generate_eda_report(client_data, validation, "artifacts/validation_report.html")

                # Step 2: Clean data
                progress.progress(30)
                clean_data = merge_and_clean(client_data, None)
                clean_data["seller_declaration_notes"] = st.session_state.seller_notes
                clean_data["notes"] = st.session_state.seller_notes
                clean_data["all_sellers"] = st.session_state.sellers
                clean_data["all_buyers"] = st.session_state.buyers

                # Step 3: Generate contract from template
                status.text("מייצר חוזה...")
                progress.progress(50)
                doc = build_contract_from_template(clean_data)
                contract_path = "artifacts/contract.docx"
                doc.save(contract_path)

                # Step 4: Save submission locally (fast, in parallel with Drive prep)
                progress.progress(60)
                save_submission(client_data, st.session_state.uploaded_docs)

                # Step 5: Upload to Google Drive
                status.text("מעלה לגוגל דרייב...")
                progress.progress(70)
                drive_result = upload_submission_to_drive(
                    client_data,
                    st.session_state.uploaded_docs,
                    artifacts_dir="artifacts",
                )
                folder_link = drive_result["folder_link"]

                # Step 6: Send email notification (only if configured)
                progress.progress(95)
                try:
                    send_notification_email(
                        seller_name=seller_name,
                        property_address=property_address,
                        drive_folder_link=folder_link,
                    )
                except Exception:
                    pass  # Email is optional, don't block on errors

                progress.progress(100)
                status.text("")

                st.session_state.sent = True
                st.success("הפרטים נשלחו לעורך הדין בהצלחה!")
                st.balloons()

                st.markdown("---")
                st.markdown(f"📎 **מסמכים שהועלו:** {drive_result['files_uploaded']}")

            except Exception as e:
                st.error(f"שגיאה: {str(e)}")
                st.info("נסה שוב או פנה לעורך הדין שלך")
